"""
Notes Exporter Module
Handles exporting generated notes to various formats (.docx, .txt).
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Union
from datetime import datetime
import io

# Document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)

class NotesExporter:
    """
    Handles exporting generated notes to various formats with customizable styling.
    """
    
    def __init__(self):
        """Initialize the exporter."""
        self.supported_formats = ['.docx', '.txt', '.md']
        
    def export_notes(self, notes_data: Dict, format_type: str = '.docx', 
                    filename: str = None) -> bytes:
        """
        Export notes to the specified format.
        
        Args:
            notes_data: Dictionary containing summary, keywords, chapters, and metadata
            format_type: Export format ('.docx', '.txt', '.md')
            filename: Optional filename (without extension)
            
        Returns:
            Exported file content as bytes
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}. Supported: {self.supported_formats}")
        
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"smart_notes_{timestamp}"
        
        try:
            if format_type == '.docx':
                return self._export_to_docx(notes_data, filename)
            elif format_type == '.txt':
                return self._export_to_txt(notes_data, filename)
            elif format_type == '.md':
                return self._export_to_markdown(notes_data, filename)
                
        except Exception as e:
            logger.error(f"Export failed: {str(e)}")
            raise Exception(f"Failed to export notes: {str(e)}")
    
    def _export_to_docx(self, notes_data: Dict, filename: str) -> bytes:
        """
        Export notes to Microsoft Word document format.
        
        Args:
            notes_data: Notes data dictionary
            filename: Base filename
            
        Returns:
            DOCX file content as bytes
        """
        # Create new document
        doc = Document()
        
        # Configure styles
        self._setup_docx_styles(doc)
        
        # Add title
        title = doc.add_heading('Smart Notes Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        self._add_metadata_section_docx(doc, notes_data.get('metadata', {}))
        
        # Add summary section
        if notes_data.get('summary'):
            self._add_summary_section_docx(doc, notes_data['summary'])
        
        # Add keywords section
        if notes_data.get('keywords'):
            self._add_keywords_section_docx(doc, notes_data['keywords'])
        
        # Add chapters section
        if notes_data.get('chapters'):
            self._add_chapters_section_docx(doc, notes_data['chapters'])
        
        # Add footer
        self._add_footer_docx(doc)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return file_stream.getvalue()
    
    def _export_to_txt(self, notes_data: Dict, filename: str) -> bytes:
        """
        Export notes to plain text format.
        
        Args:
            notes_data: Notes data dictionary
            filename: Base filename
            
        Returns:
            TXT file content as bytes
        """
        lines = []
        
        # Title
        lines.append("=" * 60)
        lines.append("SMART NOTES SUMMARY")
        lines.append("=" * 60)
        lines.append("")
        
        # Metadata
        metadata = notes_data.get('metadata', {})
        if metadata:
            lines.append("DOCUMENT INFORMATION")
            lines.append("-" * 40)
            lines.append(f"Title: {metadata.get('title', 'Unknown')}")
            lines.append(f"Author: {metadata.get('author', 'Unknown')}")
            lines.append(f"Pages: {metadata.get('pages', 'Unknown')}")
            lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            lines.append("")
        
        # Summary
        if notes_data.get('summary'):
            lines.append("SUMMARY")
            lines.append("-" * 40)
            lines.append(notes_data['summary'])
            lines.append("")
        
        # Keywords
        if notes_data.get('keywords'):
            lines.append("KEY TERMS")
            lines.append("-" * 40)
            keywords_text = ", ".join(notes_data['keywords'])
            lines.append(keywords_text)
            lines.append("")
        
        # Chapters
        if notes_data.get('chapters'):
            lines.append("CHAPTER BREAKDOWN")
            lines.append("-" * 40)
            for i, chapter in enumerate(notes_data['chapters'], 1):
                lines.append(f"Chapter {i}: {chapter.get('title', 'Untitled')}")
                lines.append("")
                content = chapter.get('content', '')
                if content:
                    # Wrap text to 80 characters
                    wrapped_content = self._wrap_text(content, 80)
                    lines.append(wrapped_content)
                lines.append("")
        
        # Footer
        lines.append("-" * 60)
        lines.append("Generated by Smart Notes Generator")
        lines.append(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    def _export_to_markdown(self, notes_data: Dict, filename: str) -> bytes:
        """
        Export notes to Markdown format.
        
        Args:
            notes_data: Notes data dictionary
            filename: Base filename
            
        Returns:
            Markdown file content as bytes
        """
        lines = []
        
        # Title
        lines.append("# Smart Notes Summary")
        lines.append("")
        
        # Metadata
        metadata = notes_data.get('metadata', {})
        if metadata:
            lines.append("## Document Information")
            lines.append("")
            lines.append("| Property | Value |")
            lines.append("|----------|-------|")
            lines.append(f"| Title | {metadata.get('title', 'Unknown')} |")
            lines.append(f"| Author | {metadata.get('author', 'Unknown')} |")
            lines.append(f"| Pages | {metadata.get('pages', 'Unknown')} |")
            lines.append(f"| Generated | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |")
            lines.append("")
        
        # Summary
        if notes_data.get('summary'):
            lines.append("## Summary")
            lines.append("")
            lines.append(notes_data['summary'])
            lines.append("")
        
        # Keywords
        if notes_data.get('keywords'):
            lines.append("## Key Terms")
            lines.append("")
            for keyword in notes_data['keywords']:
                lines.append(f"- `{keyword}`")
            lines.append("")
        
        # Chapters
        if notes_data.get('chapters'):
            lines.append("## Chapter Breakdown")
            lines.append("")
            for i, chapter in enumerate(notes_data['chapters'], 1):
                lines.append(f"### Chapter {i}: {chapter.get('title', 'Untitled')}")
                lines.append("")
                content = chapter.get('content', '')
                if content:
                    lines.append(content)
                lines.append("")
        
        # Footer
        lines.append("---")
        lines.append("*Generated by Smart Notes Generator*")
        lines.append(f"*Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        
        content = "\n".join(lines)
        return content.encode('utf-8')
    
    def _setup_docx_styles(self, doc: Document):
        """Set up custom styles for the Word document."""
        # Get or create styles
        styles = doc.styles
        
        # Normal style adjustments
        normal_style = styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        
        # Create custom styles if they don't exist
        try:
            # Section heading style
            section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.name = 'Calibri'
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = None  # Keep default color
        except ValueError:
            # Style already exists
            pass
        
        try:
            # Keyword style
            keyword_style = styles.add_style('Keyword', WD_STYLE_TYPE.CHARACTER)
            keyword_style.font.name = 'Consolas'
            keyword_style.font.size = Pt(10)
            keyword_style.font.italic = True
        except ValueError:
            # Style already exists
            pass
    
    def _add_metadata_section_docx(self, doc: Document, metadata: Dict):
        """Add metadata section to Word document."""
        if not metadata:
            return
        
        # Add section heading
        doc.add_heading('Document Information', level=1)
        
        # Create table for metadata
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Property'
        hdr_cells[1].text = 'Value'
        
        # Add metadata rows
        metadata_items = [
            ('Title', metadata.get('title', 'Unknown')),
            ('Author', metadata.get('author', 'Unknown')),
            ('Subject', metadata.get('subject', 'Unknown')),
            ('Pages', str(metadata.get('pages', 'Unknown'))),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for prop, value in metadata_items:
            row_cells = table.add_row().cells
            row_cells[0].text = prop
            row_cells[1].text = value
        
        doc.add_paragraph()  # Add spacing
    
    def _add_summary_section_docx(self, doc: Document, summary: str):
        """Add summary section to Word document."""
        doc.add_heading('Summary', level=1)
        summary_para = doc.add_paragraph(summary)
        summary_para.style = 'Normal'
        doc.add_paragraph()  # Add spacing
    
    def _add_keywords_section_docx(self, doc: Document, keywords: List[str]):
        """Add keywords section to Word document."""
        doc.add_heading('Key Terms', level=1)
        
        # Add keywords as a formatted paragraph
        keywords_para = doc.add_paragraph()
        for i, keyword in enumerate(keywords):
            if i > 0:
                keywords_para.add_run(', ')
            run = keywords_para.add_run(keyword)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.italic = True
        
        doc.add_paragraph()  # Add spacing
    
    def _add_chapters_section_docx(self, doc: Document, chapters: List[Dict]):
        """Add chapters section to Word document."""
        doc.add_heading('Chapter Breakdown', level=1)
        
        for i, chapter in enumerate(chapters, 1):
            # Chapter title
            chapter_title = f"Chapter {i}: {chapter.get('title', 'Untitled')}"
            doc.add_heading(chapter_title, level=2)
            
            # Chapter content
            content = chapter.get('content', '')
            if content:
                doc.add_paragraph(content)
            
            doc.add_paragraph()  # Add spacing between chapters
    
    def _add_footer_docx(self, doc: Document):
        """Add footer to Word document."""
        # Add page break before footer
        doc.add_page_break()
        
        # Footer information
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_run = footer_para.add_run("Generated by Smart Notes Generator")
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        
        footer_para.add_run(f"\nExported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    def _wrap_text(self, text: str, width: int = 80) -> str:
        """
        Wrap text to specified width.
        
        Args:
            text: Text to wrap
            width: Maximum line width
            
        Returns:
            Wrapped text
        """
        words = text.split()
        lines = []
        current_line = []
        current_length = 0
        
        for word in words:
            if current_length + len(word) + 1 <= width:
                current_line.append(word)
                current_length += len(word) + 1
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
                current_length = len(word)
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return '\n'.join(lines)
    
    def save_to_file(self, content: bytes, filepath: Union[str, Path], format_type: str):
        """
        Save exported content to file.
        
        Args:
            content: File content as bytes
            filepath: Path to save file
            format_type: File format extension
        """
        filepath = Path(filepath)
        
        # Ensure correct extension
        if not filepath.suffix:
            filepath = filepath.with_suffix(format_type)
        
        try:
            with open(filepath, 'wb') as f:
                f.write(content)
            logger.info(f"Notes exported successfully to {filepath}")
        except Exception as e:
            logger.error(f"Failed to save file {filepath}: {str(e)}")
            raise Exception(f"Failed to save exported file: {str(e)}")
    
    def get_export_preview(self, notes_data: Dict, format_type: str = '.txt', 
                          max_length: int = 500) -> str:
        """
        Generate a preview of the exported content.
        
        Args:
            notes_data: Notes data dictionary
            format_type: Export format
            max_length: Maximum preview length in characters
            
        Returns:
            Preview text
        """
        try:
            if format_type == '.txt':
                content = self._export_to_txt(notes_data, "preview")
                preview = content.decode('utf-8')[:max_length]
                if len(content.decode('utf-8')) > max_length:
                    preview += "..."
                return preview
            elif format_type == '.md':
                content = self._export_to_markdown(notes_data, "preview")
                preview = content.decode('utf-8')[:max_length]
                if len(content.decode('utf-8')) > max_length:
                    preview += "..."
                return preview
            else:
                return "Preview not available for this format."
        except Exception as e:
            return f"Preview generation failed: {str(e)}"
